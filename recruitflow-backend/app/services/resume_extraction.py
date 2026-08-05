"""Resume text extraction — turn raw PDF/DOCX bytes into plain text.

This is the deterministic, non-AI front half of Module 4. The Resume Parser
Agent only ever sees *text*, never the raw file, so all format-specific logic
lives here:

  * **PDF**  → PyMuPDF (fast, high-fidelity). Falls back to pdfplumber when
    PyMuPDF yields little/no text (e.g. some column layouts), and only then,
    because pdfplumber is noticeably slower.
  * **DOCX** → python-docx, pulling both paragraph runs and table cells so
    resumes laid out in tables don't lose their content.

Extraction failures raise :class:`ResumeExtractionError` so the caller (the
orchestrator runner) can log a failed ``agent_run`` instead of crashing the
background task. Scanned/image-only PDFs legitimately produce no text; we treat
"no extractable text" as an error since there's nothing for the agent to parse.
"""
from __future__ import annotations

import io
from pathlib import Path


class ResumeExtractionError(Exception):
    """Raised when a resume's text cannot be extracted."""


# Minimum characters for extracted text to be considered usable. Below this we
# assume the primary extractor failed (e.g. scanned PDF) and either fall back or
# raise, rather than handing the agent a near-empty string.
_MIN_USABLE_CHARS = 20


def _clean(text: str) -> str:
    """Normalise whitespace without destroying line structure.

    Resume structure (one item per line) is a strong signal for the parser, so
    we keep newlines but collapse runs of blank lines and trailing spaces.
    """
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    out: list[str] = []
    blank = 0
    for line in lines:
        if line.strip():
            blank = 0
            out.append(line)
        else:
            blank += 1
            if blank <= 1:  # keep single blank separators, drop the rest
                out.append("")
    return "\n".join(out).strip()


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF, PyMuPDF first with a pdfplumber fallback."""
    text = ""
    try:
        import fitz  # PyMuPDF

        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            text = "\n".join(page.get_text("text") for page in doc)
    except Exception as exc:  # noqa: BLE001 — any PyMuPDF failure → try fallback
        text = ""
        pymupdf_error: Exception | None = exc
    else:
        pymupdf_error = None

    # Fall back to pdfplumber only when PyMuPDF gave us little/nothing, since
    # pdfplumber is slower. This rescues some column/table layouts.
    if len(text.strip()) < _MIN_USABLE_CHARS:
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as exc:  # noqa: BLE001
            if pymupdf_error is not None:
                raise ResumeExtractionError(
                    f"Both PDF extractors failed: PyMuPDF={pymupdf_error!r}, "
                    f"pdfplumber={exc!r}"
                ) from exc
            # PyMuPDF succeeded-but-thin and pdfplumber errored; keep what we had.

    return _clean(text)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX, including table cells (not just paragraphs)."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001
        raise ResumeExtractionError(f"Could not open DOCX: {exc!r}") from exc

    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(c for c in cells if c))

    return _clean("\n".join(parts))


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a resume file.

    ``filename`` is used only to pick the extractor by extension; the intake
    layer has already validated it is a PDF or DOCX. Raises
    :class:`ResumeExtractionError` for unsupported types or when no usable text
    can be pulled out (e.g. scanned/image-only PDFs).
    """
    if not file_bytes:
        raise ResumeExtractionError("Resume file is empty")

    ext = Path(filename or "").suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(file_bytes)
    elif ext == ".docx":
        text = _extract_docx(file_bytes)
    else:
        raise ResumeExtractionError(
            f"Unsupported resume type {ext!r}; expected .pdf or .docx"
        )

    if len(text.strip()) < _MIN_USABLE_CHARS:
        raise ResumeExtractionError(
            "No extractable text found — the file may be scanned or image-only"
        )

    return text
